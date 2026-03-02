"""ingest_document ARQ task — processes user-uploaded files into knowledge chunks.

Supports PDF, DOCX, TXT, and MD. After chunking and embedding, runs
refine_intent to detect domain shifts and update structured_intent.
"""

import base64
import io
import json
from datetime import UTC, datetime

import structlog

from app.core.cache_invalidation import invalidate_project_caches
from app.core.db import get_db_pool
from app.core.redis import get_redis
from app.kb.ingester import RawDocument, ingest_documents
from app.kb.intent_extractor import refine_intent
from app.repositories import project as project_repo

logger = structlog.get_logger(__name__)


def _coerce_intent_mapping(value: object) -> dict:
    """Normalize structured intent payload into a dictionary."""
    if isinstance(value, dict):
        return dict(value)
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return {}
        try:
            loaded = json.loads(text)
        except json.JSONDecodeError:
            return {}
        return dict(loaded) if isinstance(loaded, dict) else {}
    return {}


def _extract_text_from_bytes(content_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from uploaded file bytes based on file extension.

    Returns extracted text string. Returns "" on parse failure so the caller
    can skip the document rather than crashing the worker.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    try:
        if ext == "pdf":
            import pypdf

            reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages).strip()

        elif ext == "docx":
            import docx

            doc = docx.Document(io.BytesIO(content_bytes))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

        elif ext in ("txt", "md"):
            return content_bytes.decode("utf-8", errors="replace").strip()

        else:
            logger.warning("ingest_document.unsupported_ext", ext=ext, filename=filename)
            return ""

    except Exception:
        logger.exception("ingest_document.extract_error", filename=filename, ext=ext)
        return ""


async def ingest_document(
    ctx: dict,
    upload_id: str,
    project_id: str,
    user_id: str,
    staging_key: str,
    filename: str,
) -> None:
    """
    Process an uploaded document: extract text, chunk, embed, upsert to KB,
    then refine structured_intent if new content shifts the project domain.

    upload_id is used as source_id so re-uploading the same file is idempotent
    (ON CONFLICT DO NOTHING in the ingester).
    """
    content_bytes = b""
    size_bytes = 0
    redis = await get_redis()
    try:
        payload = await redis.get(staging_key)
        if isinstance(payload, str):
            # Payload is base64-encoded (stored by kb.py upload handler)
            content_bytes = base64.b64decode(payload)
        elif isinstance(payload, bytes):
            # Fallback for legacy pre-base64 entries
            content_bytes = payload
        if not content_bytes:
            logger.warning(
                "ingest_document.staging_payload_missing",
                upload_id=upload_id,
                staging_key=staging_key,
            )
            return
        size_bytes = len(content_bytes)
    except Exception:
        logger.exception(
            "ingest_document.staging_payload_read_error",
            upload_id=upload_id,
            staging_key=staging_key,
        )
        return

    logger.info(
        "ingest_document.start",
        upload_id=upload_id,
        project_id=project_id,
        filename=filename,
        size_bytes=size_bytes,
        staging_key=staging_key,
    )

    try:
        # Step 1: Extract text
        text = _extract_text_from_bytes(content_bytes, filename)
        if not text:
            logger.warning("ingest_document.empty_text", upload_id=upload_id, filename=filename)
            return

        # Step 2: Build RawDocument and ingest
        documents = [
            RawDocument(
                project_id=project_id,
                user_id=user_id,
                source="upload",
                source_id=upload_id,
                title=filename,
                content=text,
                metadata={"filename": filename, "upload_id": upload_id},
            )
        ]
        inserted = await ingest_documents(documents)
        logger.info("ingest_document.ingested", upload_id=upload_id, chunks_inserted=inserted)

        # Invalidate caches so users see new content immediately
        await invalidate_project_caches(project_id)

        # Step 3: Refine intent — collect short summaries from the new chunks
        pool = await get_db_pool()
        async with pool.acquire() as conn:
            await conn.execute("SELECT set_config('app.accessible_projects', $1, true)", project_id)
            result = await conn.fetch(
                """
                WITH recent_chunks AS (
                    SELECT content FROM knowledge_chunks
                    WHERE project_id = $1::uuid AND source_id LIKE $2
                    ORDER BY created_at DESC
                    LIMIT 20
                ),
                chunk_count AS (
                    SELECT COUNT(*) as total FROM knowledge_chunks
                    WHERE project_id = $1::uuid
                )
                SELECT
                    recent_chunks.content,
                    chunk_count.total
                FROM recent_chunks
                CROSS JOIN chunk_count
                """,
                project_id,
                f"{upload_id}%",
            )
        chunk_summaries = [row["content"][:200] for row in result]
        total = result[0]["total"] if result else 0

        if chunk_summaries:
            project = await project_repo.fetch_project(pool, project_id, user_id)
            if project:
                existing_intent = _coerce_intent_mapping(project.get("structured_intent"))
                try:
                    refined = await refine_intent(existing_intent, chunk_summaries)
                    if refined != existing_intent:
                        await project_repo.update_project_intent(pool, project_id, refined)
                        logger.info(
                            "ingest_document.intent_refined",
                            project_id=project_id,
                            new_domain=refined.get("domain"),
                        )
                except Exception:
                    logger.exception(
                        "ingest_document.intent_refine_failed",
                        project_id=project_id,
                        upload_id=upload_id,
                    )

        # Step 4: Update kb_chunk_count
        await project_repo.update_project_kb_stats(
            pool, project_id, int(total or 0), datetime.now(UTC)
        )

        logger.info(
            "ingest_document.done",
            upload_id=upload_id,
            project_id=project_id,
            total_chunks=total,
        )
    finally:
        try:
            await redis.delete(staging_key)
        except Exception:
            logger.warning(
                "ingest_document.staging_cleanup_failed",
                upload_id=upload_id,
                staging_key=staging_key,
            )
